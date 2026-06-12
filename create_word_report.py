import sys
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("python-docx is not installed. Please install it first.")
    sys.exit(1)

def create_report():
    document = Document()

    # Title
    title = document.add_heading('Ava: Real-Time Voice Assistant Pipeline', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Introduction
    document.add_heading('1. Introduction', level=1)
    p = document.add_paragraph(
        'Ava is a real-time, highly capable voice assistant pipeline designed for low latency and seamless interaction. '
        'The architecture is built on a multithreaded node-based system that allows parallel processing of audio input, '
        'speech-to-text (STT) transcription, conversational memory retrieval, large language model (LLM) text generation, '
        'and text-to-speech (TTS) synthesis.'
    )

    # Architecture Overview
    document.add_heading('2. System Architecture', level=1)
    p = document.add_paragraph('The pipeline is broken down into several independent nodes communicating via thread-safe queues. '
        'This allows the system to process incoming audio streams concurrently while generating responses, significantly reducing the perceived latency.'
    )
    
    document.add_paragraph('Audio Input Node (audio.py)', style='List Bullet')
    document.add_paragraph('Speech-to-Text Node (stt.py)', style='List Bullet')
    document.add_paragraph('Memory Node (memory.py)', style='List Bullet')
    document.add_paragraph('LLM Node (llm.py)', style='List Bullet')
    document.add_paragraph('Text-to-Speech Node (tts.py)', style='List Bullet')

    # Core Components Details
    document.add_heading('3. Core Components', level=1)
    
    document.add_heading('3.1. Speech-to-Text (STT)', level=2)
    document.add_paragraph(
        'The STT node leverages the faster-whisper library, specifically configured to run on the CPU with int8 precision. '
        'It listens to chunks of audio from the audio input queue and processes them using a small Whisper model, optimized for accuracy and fast execution.'
    )

    document.add_heading('3.2. Conversational Memory', level=2)
    document.add_paragraph(
        'The Memory node is powered by ChromaDB, a persistent vector database. '
        'Using the SentenceTransformers embedding function (all-MiniLM-L6-v2), it stores transcriptions and retrieves relevant past context (up to 3 recent results) '
        'to enrich the current prompt for the LLM. This allows Ava to maintain context over long conversations.'
    )

    document.add_heading('3.3. Language Model (LLM)', level=2)
    document.add_paragraph(
        'The LLM node connects to a locally hosted instance of Ollama. It uses a lightweight model (llama3.2:1b) which is highly capable but small enough to run effectively on consumer hardware. '
        'The text generation is streamed, and responses are broken down into natural sentence chunks to be piped into the TTS node immediately, minimizing time-to-first-audio.'
    )

    document.add_heading('3.4. Text-to-Speech (TTS)', level=2)
    document.add_paragraph(
        'The TTS node relies on Piper TTS (en_US-amy-medium.onnx), an extremely fast and natural-sounding engine. '
        'As the LLM generates sentence chunks, the TTS node synthesizes the audio and streams it to the system output using the sounddevice library.'
    )

    # Configuration Details
    document.add_heading('4. Configuration & Setup', level=1)
    document.add_paragraph(
        'The system behavior is governed by a central config.py file which handles variables such as sample rates, chunk sizes, VAD (Voice Activity Detection) thresholds, '
        'and model endpoints. The main.py entry point gracefully orchestrates the nodes, handling interruptions seamlessly so that Ava stops speaking if the user interrupts her.'
    )

    document.add_heading('5. Conclusion', level=1)
    document.add_paragraph(
        'Ava represents a robust, locally hosted edge AI architecture. By decoupling components and employing a queue-based asynchronous design, it provides an engaging, responsive, and privacy-preserving voice interaction experience.'
    )

    file_path = 'Ava_Project_Report.docx'
    document.save(file_path)
    print(f"Report successfully saved to {file_path}")

if __name__ == '__main__':
    create_report()
